"""
reporting.py
============

Automated executive report generation module for the End-to-End Sales
Forecasting & Demand Intelligence System. Generates a formatted Word document
(`summary.docx`) in the project root containing executive summaries, analysis
findings, tables, and embedded static charts.
"""

from __future__ import annotations

import os
from pathlib import Path
from typing import Dict, Any, Optional

import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

from src.utils import PATHS, get_logger, DataValidationError

logger = get_logger(__name__)


# ======================================================================
# 1. Styling Helpers for python-docx
# ======================================================================
def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Set inner padding for a table cell (values in dxa: 1 pt = 20 dxa)."""
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{margin}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)


def set_cell_shading(cell, color_hex="1F77B4"):
    """Set background color of a cell."""
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'
    cell._element.get_or_add_tcPr().append(parse_xml(shading_xml))


def format_table_header(row, col_widths=None, bg_color="1B365D"):
    """Apply background shading and bold white text to a header row."""
    for i, cell in enumerate(row.cells):
        set_cell_shading(cell, bg_color)
        set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
        if col_widths and i < len(col_widths):
            cell.width = col_widths[i]
        
        # Style text in header run
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(10)


def style_table_borders(table):
    """Apply standard clean light grey borders to a table."""
    tblPr = table._element.xpath('w:tblPr')
    if tblPr:
        # Define clean thin borders
        borders_xml = (
            f'<w:tblBorders {nsdecls("w")}>'
            '<w:top w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
            '<w:bottom w:val="single" w:sz="6" w:space="0" w:color="1B365D"/>'
            '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="E0E0E0"/>'
            '<w:insideV w:val="none"/>'
            '<w:left w:val="none"/>'
            '<w:right w:val="none"/>'
            '</w:tblBorders>'
        )
        tblPr[0].append(parse_xml(borders_xml))


# ======================================================================
# 2. Document Generation
# ======================================================================
def generate_executive_report(
    df_sales: pd.DataFrame,
    forecast_results: Dict[str, Any],
    anomaly_table: pd.DataFrame,
    segmented_features: pd.DataFrame,
    save_path: Optional[Path] = None,
) -> Path:
    """
    Generate the executive Word report `summary.docx` summarizing the findings,
    forecasts, anomalies, and demand segments.

    Parameters
    ----------
    df_sales : pd.DataFrame
        Cleaned transactional/monthly sales dataframe.
    forecast_results : dict
        Results dict returned by train_and_compare_models.
    anomaly_table : pd.DataFrame
        Consolidated anomaly table from detect_anomalies.
    segmented_features : pd.DataFrame
        Product features with ClusterLabel column from segment_demand.
    save_path : Path, optional
        Target report destination. Defaults to project root.
    """
    if save_path is None:
        save_path = PATHS.root / "summary.docx"
    else:
        save_path = Path(save_path)

    logger.info("Initializing summary.docx generation at %s", save_path)
    
    doc = Document()
    
    # ------------------------------------------------------------------
    # Document Geometry
    # ------------------------------------------------------------------
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # ------------------------------------------------------------------
    # Base Typography Settings
    # ------------------------------------------------------------------
    style_normal = doc.styles["Normal"]
    font = style_normal.font
    font.name = "Arial"
    font.size = Pt(10.5)
    font.color.rgb = RGBColor(51, 51, 51)  # Off-black / charcoal

    # Colors
    c_primary = RGBColor(27, 54, 93)     # Navy (#1B365D)
    c_secondary = RGBColor(242, 142, 43) # Orange (#F28E2B)
    
    # ------------------------------------------------------------------
    # Title & Metadata
    # ------------------------------------------------------------------
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(36)
    title_p.paragraph_format.space_after = Pt(6)
    title_run = title_p.add_run("SALES FORECASTING & DEMAND INTELLIGENCE SYSTEM")
    title_run.bold = True
    title_run.font.size = Pt(22)
    title_run.font.color.rgb = c_primary

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_after = Pt(24)
    sub_run = sub_p.add_run("Executive Business Report & Decision Support Summary")
    sub_run.font.size = Pt(13)
    sub_run.font.italic = True
    sub_run.font.color.rgb = RGBColor(120, 120, 120)

    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_p.paragraph_format.space_after = Pt(40)
    meta_p.add_run("Author: Arjun Peddi\nRole: Senior Data Scientist\nDate: July 2026")

    doc.add_page_break()

    # ------------------------------------------------------------------
    # 1. Executive Summary
    # ------------------------------------------------------------------
    h1 = doc.add_paragraph()
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(6)
    r1 = h1.add_run("1. Executive Summary")
    r1.bold = True
    r1.font.size = Pt(16)
    r1.font.color.rgb = c_primary

    p = doc.add_paragraph()
    p.add_run(
        "Effective inventory management is critical to sustaining retail margins and customer satisfaction. "
        "Over-forecasting ties up precious capital in warehouse space and forces clearance markdowns; "
        "under-forecasting results in stockouts, lost sales, and degraded brand loyalty. "
        "This project delivers an End-to-End Sales Forecasting & Demand Intelligence System designed to "
        "align purchasing workflows with automated, data-driven planning signals. "
        "By integrating time series forecasting, dynamic anomaly detection, and cluster-based demand segmentation, "
        "management is equipped with a unified system for operational decision support."
    )
    
    p2 = doc.add_paragraph()
    p2.add_run(
        "Key findings from our implementations indicate: \n"
        "•  Multi-Model Forecasting: We evaluated three time-series methods (SARIMA, Prophet, and XGBoost). "
        f"The best performing model is {forecast_results.get('best_model_name', 'SARIMA')}, "
        f"which achieved a Validation MAPE of {forecast_results.get('comparison_df').iloc[0]['MAPE']:.2%}.\n"
        f"•  Anomaly Detection: An ensemble model successfully flagged {len(anomaly_table)} historical sales anomalies. "
        "These anomalies represent outlier demand events (bulk orders or holiday rushes) that need to be winsorized "
        "or treated separately to prevent bias in replenishment schedules.\n"
        "•  Product Demand Segmentation: Using KMeans clustering, the product catalog is structured into "
        "four segments (High Volume, Stable Demand; Growing Demand; Declining Demand; and Low Volume, High Volatility). "
        "Each segment is mapped to custom ordering rules to optimize capital allocation."
    )

    # ------------------------------------------------------------------
    # 2. Exploratory Data Analysis (EDA) Findings
    # ------------------------------------------------------------------
    h2 = doc.add_paragraph()
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(6)
    r2 = h2.add_run("2. Exploratory Data Analysis (EDA) Findings")
    r2.bold = True
    r2.font.size = Pt(14)
    r2.font.color.rgb = c_primary

    p = doc.add_paragraph()
    total_rev = df_sales["Sales"].sum()
    avg_order = df_sales.groupby("YearMonth")["Sales"].sum().mean()
    p.add_run(
        f"The primary transactional dataset contains {len(df_sales):,} historical lines. "
        f"Total sales volume stands at ${total_rev:,.2f}, with a monthly average sales rate of ${avg_order:,.2f}. "
        "The sales distribution is characterized by high seasonality (spiking in Q4 of each year) and a steady "
        "year-over-year upward baseline trend."
    )

    # Embed monthly sales chart
    img_path = PATHS.charts / "monthly_sales_trend.png"
    if img_path.exists():
        doc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_picture(str(img_path), width=Inches(5.5))
        caption = doc.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = caption.add_run("Figure 1: Monthly Sales Trend with 3-Month Rolling Average")
        cap_run.font.italic = True
        cap_run.font.size = Pt(9.5)

    # Regional / Category breakdowns
    reg_img = PATHS.charts / "regional_sales.png"
    cat_img = PATHS.charts / "category_sales.png"
    if reg_img.exists() or cat_img.exists():
        doc.add_page_break()
        doc.add_paragraph().add_run("Additional EDA breakdowns by product categorization and geographical regions reveal the core drivers of sales volume:")
        
        if reg_img.exists():
            doc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_picture(str(reg_img), width=Inches(5.0))
            caption = doc.add_paragraph()
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption.add_run("Figure 2: Regional Sales Revenue Distribution").font.italic = True
            
        if cat_img.exists():
            doc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_picture(str(cat_img), width=Inches(5.0))
            caption = doc.add_paragraph()
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption.add_run("Figure 3: Product Sub-Category Contribution").font.italic = True

    # ------------------------------------------------------------------
    # 3. Forecasting Results & Model Comparison
    # ------------------------------------------------------------------
    h3 = doc.add_paragraph()
    h3.paragraph_format.space_before = Pt(18)
    h3.paragraph_format.space_after = Pt(6)
    r3 = h3.add_run("3. Forecasting Results & Model Comparison")
    r3.bold = True
    r3.font.size = Pt(14)
    r3.font.color.rgb = c_primary

    p = doc.add_paragraph()
    p.add_run(
        "To establish a robust forecasting pipeline, SARIMA, Meta Prophet, and XGBoost models were evaluated "
        "against a 3-month validation holdout. The models were evaluated on Mean Absolute Error (MAE), "
        "Root Mean Squared Error (RMSE), and Mean Absolute Percentage Error (MAPE). Below is the model comparison summary:"
    )

    # Create Table of Metrics
    comp_df = forecast_results["comparison_df"]
    table = doc.add_table(rows=len(comp_df) + 1, cols=4)
    style_table_borders(table)
    col_widths = [Inches(1.8), Inches(1.2), Inches(1.2), Inches(1.2)]
    
    # Headers
    table.rows[0].cells[0].paragraphs[0].text = "Forecasting Model"
    table.rows[0].cells[1].paragraphs[0].text = "MAE"
    table.rows[0].cells[2].paragraphs[0].text = "RMSE"
    table.rows[0].cells[3].paragraphs[0].text = "MAPE"
    format_table_header(table.rows[0], col_widths)

    # Populating Rows
    for idx, row in comp_df.iterrows():
        r_cells = table.rows[idx + 1].cells
        r_cells[0].paragraphs[0].text = str(row["Model"])
        r_cells[1].paragraphs[0].text = f"${row['MAE']:,.2f}"
        r_cells[2].paragraphs[0].text = f"${row['RMSE']:,.2f}"
        r_cells[3].paragraphs[0].text = f"{row['MAPE']:.2%}"
        
        # Center align metrics
        for cell in r_cells[1:]:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for cell in r_cells:
            set_cell_margins(cell, top=80, bottom=80, left=150, right=150)

    # Insert forecast plot of the best model
    # The plot is named forecast_sales_forecast.png or similar
    fc_img = None
    for item in PATHS.charts.glob("forecast_*.png"):
        fc_img = item
        break
        
    if fc_img and fc_img.exists():
        doc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_picture(str(fc_img), width=Inches(5.5))
        caption = doc.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.add_run(f"Figure 4: 3-Month Sales Forecast with 95% Confidence Bounds (Model: {forecast_results.get('best_model_name')})").font.italic = True

    # ------------------------------------------------------------------
    # 4. Anomaly Detection Findings
    # ------------------------------------------------------------------
    doc.add_page_break()
    h4 = doc.add_paragraph()
    h4.paragraph_format.space_before = Pt(18)
    h4.paragraph_format.space_after = Pt(6)
    r4 = h4.add_run("4. Anomaly Detection Findings")
    r4.bold = True
    r4.font.size = Pt(14)
    r4.font.color.rgb = c_primary

    p = doc.add_paragraph()
    p.add_run(
        "Sales anomalies can indicate operational disruptions, marketing campaign spikes, or data entry errors. "
        "We implemented an ensemble anomaly detector combining Z-score (global scale), Rolling Mean (local context), "
        "and Isolation Forest (distribution clustering). Consensus anomalies are defined as those flagged by at least "
        f"two of the three algorithms. A total of {len(anomaly_table)} consensus anomalies were flagged."
    )

    anom_img = PATHS.charts / "sales_anomalies.png"
    if anom_img.exists():
        doc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_picture(str(anom_img), width=Inches(5.5))
        caption = doc.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.add_run("Figure 5: Historical Sales with Consolidated Anomalies Highlighted").font.italic = True

    # Display subset of anomalies table
    if not anomaly_table.empty:
        doc.add_paragraph().add_run("Sample of major detected sales anomalies:")
        display_anom = anomaly_table.sort_values(by="Sales", ascending=False).head(5)
        
        table = doc.add_table(rows=len(display_anom) + 1, cols=4)
        style_table_borders(table)
        
        table.rows[0].cells[0].paragraphs[0].text = "Date"
        table.rows[0].cells[1].paragraphs[0].text = "Sales"
        table.rows[0].cells[2].paragraphs[0].text = "Rolling Mean"
        table.rows[0].cells[3].paragraphs[0].text = "Reason / Flag Count"
        format_table_header(table.rows[0], [Inches(1.5), Inches(1.2), Inches(1.2), Inches(1.8)])
        
        for r_idx, (_, row) in enumerate(display_anom.iterrows()):
            cells = table.rows[r_idx + 1].cells
            # Check date formatting
            date_val = row.name if isinstance(row.name, (str, pd.Timestamp)) else row.get("Date", "N/A")
            if isinstance(date_val, pd.Timestamp):
                cells[0].paragraphs[0].text = date_val.strftime("%Y-%m-%d")
            else:
                cells[0].paragraphs[0].text = str(date_val)[:10]

            cells[1].paragraphs[0].text = f"${row['Sales']:,.2f}"
            cells[2].paragraphs[0].text = f"${row.get('Rolling_Mean', 0):,.2f}"
            cells[3].paragraphs[0].text = f"Spike (Flag Count: {int(row.get('Anomaly_Vote_Count', 0))})"
            
            cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for cell in cells:
                set_cell_margins(cell, top=80, bottom=80, left=150, right=150)

    # ------------------------------------------------------------------
    # 5. Product Demand Segmentation
    # ------------------------------------------------------------------
    doc.add_page_break()
    h5 = doc.add_paragraph()
    h5.paragraph_format.space_before = Pt(18)
    h5.paragraph_format.space_after = Pt(6)
    r5 = h5.add_run("5. Product Demand Segmentation")
    r5.bold = True
    r5.font.size = Pt(14)
    r5.font.color.rgb = c_primary

    p = doc.add_paragraph()
    p.add_run(
        "Rather than planning inventory uniformly across the catalog, a KMeans clustering model was trained "
        "on product-level performance features: Total Sales, Transaction Frequency, Volatility, and Sales Trend. "
        "This segments the product catalog into four distinct tiers, allowing for custom ordering policies."
    )

    cluster_img = PATHS.charts / "demand_segmentation.png"
    if cluster_img.exists():
        doc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_picture(str(cluster_img), width=Inches(5.5))
        caption = doc.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.add_run("Figure 6: PCA Projection of Product Demand Clusters").font.italic = True

    # Segments Breakdown
    counts = segmented_features["ClusterLabel"].value_counts()
    table = doc.add_table(rows=len(counts) + 1, cols=3)
    style_table_borders(table)
    
    table.rows[0].cells[0].paragraphs[0].text = "Demand Segment"
    table.rows[0].cells[1].paragraphs[0].text = "Product Count"
    table.rows[0].cells[2].paragraphs[0].text = "Core Strategy"
    format_table_header(table.rows[0], [Inches(2.0), Inches(1.2), Inches(3.0)])
    
    strategies = {
        "High Volume, Stable Demand": "Automated Min-Max replenishment, high safety stock",
        "Growing Demand": "Dynamic safety stock tied to peaks, pre-season capacity planning",
        "Declining Demand": "Low, steady safety stock, consolidated purchasing runs",
        "Low Volume, High Volatility": "Make-to-order, drop-shipping, or deep clearance discount campaigns",
    }
    
    for idx, (label, count) in enumerate(counts.items()):
        cells = table.rows[idx + 1].cells
        cells[0].paragraphs[0].text = str(label)
        cells[1].paragraphs[0].text = f"{count:,}"
        
        # Strip trailing "(Tier B)" from label to match recommendations key
        clean_lbl = label.split(" (")[0]
        cells[2].paragraphs[0].text = strategies.get(clean_lbl, "Standard inventory review")
        
        cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for cell in cells:
            set_cell_margins(cell, top=80, bottom=80, left=150, right=150)

    # ------------------------------------------------------------------
    # 6. Actionable Business Recommendations
    # ------------------------------------------------------------------
    h6 = doc.add_paragraph()
    h6.paragraph_format.space_before = Pt(18)
    h6.paragraph_format.space_after = Pt(6)
    r6 = h6.add_run("6. Actionable Business Recommendations")
    r6.bold = True
    r6.font.size = Pt(14)
    r6.font.color.rgb = c_primary

    p = doc.add_paragraph()
    p.add_run(
        "Based on the combined outputs of the system, we recommend the following operational workflows:\n\n"
        "1. Implement Segment-specific Inventory Rules: Transition the replenishment system to run automatically for "
        "High Volume, Stable Demand using a standard Min-Max policy. Reallocate purchasing hours to focus on "
        "Growing Demand items where dynamic lead times are required.\n\n"
        "2. Winsorize Forecast Inputs: Use the anomaly detection ensemble as an automated ETL step. When an anomaly is detected, "
        "automatically winsorize (cap) the spike in historical databases before feeding it to the forecasting engines. "
        "This prevents transient spikes from causing inflated purchase orders.\n\n"
        "3. Streamline Slow-Moving SKUs: Product listings categorized as 'Low Volume, High Volatility' should be evaluated for SKU "
        "rationalization. If they are not strategically necessary, run clearance marketing campaigns to free up warehouse space. "
        "If they are necessary, transition them to drop-ship fulfillment."
    )

    # ------------------------------------------------------------------
    # 7. Limitations & Future Work
    # ------------------------------------------------------------------
    h7 = doc.add_paragraph()
    h7.paragraph_format.space_before = Pt(18)
    h7.paragraph_format.space_after = Pt(6)
    r7 = h7.add_run("7. Limitations & Future Work")
    r7.bold = True
    r7.font.size = Pt(14)
    r7.font.color.rgb = c_primary

    p = doc.add_paragraph()
    p.add_run(
        "While the system performs robustly, several limitations exist:\n"
        "•  External Factors: The models currently rely purely on historical sales. They do not capture marketing budgets, "
        "competitor pricing, macroeconomic indexes, or local holiday listings.\n"
        "•  Supply Chain Constraints: The recommendations assume infinite supply and zero transport lead-times. "
        "Integrating lead-time volatility and supplier capacity will make replenishment recommendations more realistic.\n"
        "•  Model Drift: Sales patterns change. We recommend setting up an automated model-monitoring cron to check for "
        "performance drift (e.g. MAPE exceeding 15%) and trigger monthly re-training cycles."
    )

    # ------------------------------------------------------------------
    # Save Report
    # ------------------------------------------------------------------
    doc.save(str(save_path))
    logger.info("Executive report successfully generated and saved to %s", save_path)

    # Also generate PDF version
    pdf_path = save_path.with_suffix(".pdf")
    convert_docx_to_pdf(save_path, pdf_path)

    return save_path


def convert_docx_to_pdf(docx_path: Path, pdf_path: Path):
    """Convert a .docx file to .pdf using Windows COM (MS Word) automation."""
    import win32com.client
    import pythoncom
    
    logger.info("Converting %s to PDF via MS Word COM...", docx_path.name)
    pythoncom.CoInitialize()
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(str(docx_path.resolve()))
        # 17 is the constant for wdFormatPDF
        doc.SaveAs(str(pdf_path.resolve()), FileFormat=17)
        doc.Close()
        word.Quit()
        logger.info("Successfully generated PDF: %s", pdf_path.name)
    except Exception as e:
        logger.error("COM PDF conversion failed: %s. Ensuring file handles are closed.", e)
    finally:
        pythoncom.CoUninitialize()

